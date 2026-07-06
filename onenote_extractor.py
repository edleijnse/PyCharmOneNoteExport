import re
import aspose.note as asn
from docx import Document


def sanitize_xml_string(s: str) -> str:
    """Remove characters that are not allowed in XML 1.0."""
    if not s:
        return ""
    # Filter out control characters that are invalid in XML
    # Valid: #x9, #xA, #xD, [#x20-#xD7FF], [#xE000-#xFFFD], [#x10000-#x10FFFF]
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', s)


def load(one_file: str) -> asn.Document:
    return asn.Document(one_file)


def get_page_text(page: asn.Page) -> str:
    lines = []
    for rt in page.GetChildNodes(asn.RichText):
        text = rt.Text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def get_page_title(page: asn.Page) -> str:
    try:
        return page.Title.TitleText.Text.strip()

    except Exception:
        pass
    return "(untitled)"


def extract_pages(doc: asn.Document, section_name: str = "") -> list[dict]:
    """Return one dict per page with section, page name, level, and text."""
    if not section_name:
        section_name = "Section"

    results = []
    for page in doc:
        results.append({
            "section": section_name,
            "page": get_page_title(page),
            # "level": page.level,
            "text": get_page_text(page),
        })
    return results


def print_pages(pages: list[dict]) -> None:
    pad = "  "
    for entry in pages:
        # indent = pad * (entry["level"] - 1)
        indent = pad
        print("=" * 72)
        print(f"{indent}[{entry['section']}] > {entry['page']}")
        print("-" * 72)
        if entry["text"]:
            for line in entry["text"].splitlines():
                print(f"{indent}  {line}")
        else:
            print(f"{indent}  (no text content)")
        print()


def write_pages(pages: list[dict], output_file: str) -> None:
    pad = "  "
    with open(output_file, "w", encoding="utf-8") as f:
        for entry in pages:
            # indent = pad * (entry["level"] - 1)
            indent = pad
            f.write("=" * 72 + "\n")
            f.write(f"{indent}[{entry['section']}] > {entry['page']}\n")
            f.write("-" * 72 + "\n")
            if entry["text"]:
                for line in entry["text"].splitlines():
                    f.write(f"{indent}  {line}\n")
            else:
                f.write(f"{indent}  (no text content)\n")
            f.write("\n")


def write_pages_to_doc(pages: list[dict], output_file: str) -> None:
    doc = Document()
    for entry in pages:
        title = sanitize_xml_string(f"[{entry['section']}] > {entry['page']}")
        doc.add_heading(title, level=1)
        if entry["text"]:
            text = sanitize_xml_string(entry["text"])
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("(no text content)")
    doc.save(output_file)


if __name__ == "__main__":
    import os

    ONE_FILE = r"D:\OneNote\Paspoort.one"
    ONE_FILE_OUTPUT = r"D:\OneNote\Paspoort.one.text"
    ONE_FILE_OUTPUT_DOCX = r"D:\OneNote\Paspoort.one.docx"
    section_name = os.path.splitext(os.path.basename(ONE_FILE))[0]

    doc = load(ONE_FILE)
    pages = extract_pages(doc, section_name)
    print_pages(pages)
    write_pages(pages, ONE_FILE_OUTPUT)
    write_pages_to_doc(pages, ONE_FILE_OUTPUT_DOCX)
